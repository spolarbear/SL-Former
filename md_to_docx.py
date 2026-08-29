# -*- coding: utf-8 -*-
"""把论文 Markdown 转换为 Word (.docx) 文档.
用法: python md_to_docx.py
"""
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_markdown_text(doc, text):
    """解析行内粗体/代码/公式, 添加带格式的段落."""
    p = doc.add_paragraph()
    # 分割: **bold**, `code`, $math$
    tokens = re.split(r'(\*\*.+?\*\*|`[^`]+`|\$[^$]+\$)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            run = p.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            run = p.add_run(tok[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        elif tok.startswith('$') and tok.endswith('$'):
            run = p.add_run(tok)  # 保留 LaTeX 公式文本
            run.italic = True
            run.font.name = 'Cambria Math'
        else:
            p.add_run(tok)
    return p


def convert(md_path, out_path):
    doc = Document()
    # 默认字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    # 中文
    from docx.oxml.ns import qn
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # 跳过 mermaid 代码块 (流程图用文本示意)
        if stripped.startswith('```'):
            # 收集代码块内容
            block = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                block.append(lines[i].rstrip('\n'))
                i += 1
            i += 1
            if block and block[0].strip().lower() == 'mermaid':
                # mermaid 图: 以文本段落形式加入
                doc.add_paragraph('[Mermaid 流程图]')
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(block[1:]))
                run.font.name = 'Consolas'
                run.font.size = Pt(8)
            continue

        # 标题
        m = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if m:
            level = len(m.group(1))
            text = re.sub(r'[*`$]', '', m.group(2))
            h = doc.add_heading(text, level=level)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue

        # 表格
        if stripped.startswith('|'):
            rows = []
            while i < n and lines[i].strip().startswith('|'):
                rows.append(lines[i].strip())
                i += 1
            # 去掉分隔行
            rows = [r for r in rows if not re.match(r'^\|[\s\-:|]+\|$', r)]
            if rows:
                cells = [[c.strip() for c in r.strip('|').split('|')] for r in rows]
                ncol = max(len(c) for c in cells)
                table = doc.add_table(rows=len(cells), cols=ncol)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for ri, row in enumerate(cells):
                    for ci in range(ncol):
                        cell_txt = row[ci] if ci < len(row) else ''
                        cell_txt = re.sub(r'[*`$]', '', cell_txt)
                        cell = table.cell(ri, ci)
                        cell.text = cell_txt
                        for par in cell.paragraphs:
                            for run in par.runs:
                                run.font.size = Pt(9)
                            if ri == 0:
                                for run in par.runs:
                                    run.bold = True
                doc.add_paragraph()
            continue

        # 引用
        if stripped.startswith('>'):
            p = doc.add_paragraph()
            run = p.add_run(stripped.lstrip('> '))
            run.italic = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            i += 1
            continue

        # 分隔线
        if stripped in ('---', '***', '___'):
            i += 1
            continue

        # 列表
        m = re.match(r'^(\s*)[-*+]\s+(.*)', stripped)
        if m:
            p = doc.add_paragraph(style='List Bullet')
            add_markdown_text(doc, '')
            # 重新用列表样式
            p = doc.paragraphs[-1]
            p.style = doc.styles['List Bullet']
            # 清空刚加的
            p._element.getparent().remove(p._element)
            p = doc.add_paragraph(style='List Bullet')
            tokens = re.split(r'(\*\*.+?\*\*|`[^`]+`|\$[^$]+\$)', m.group(2))
            for tok in tokens:
                if not tok:
                    continue
                if tok.startswith('**') and tok.endswith('**'):
                    r = p.add_run(tok[2:-2]); r.bold = True
                elif tok.startswith('`') and tok.endswith('`'):
                    r = p.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9)
                elif tok.startswith('$') and tok.endswith('$'):
                    r = p.add_run(tok); r.italic = True
                else:
                    p.add_run(tok)
            i += 1
            continue

        # 有序列表
        m = re.match(r'^(\d+)[.)]\s+(.*)', stripped)
        if m:
            p = doc.add_paragraph(style='List Number')
            tokens = re.split(r'(\*\*.+?\*\*|`[^`]+`|\$[^$]+\$)', m.group(2))
            for tok in tokens:
                if not tok:
                    continue
                if tok.startswith('**') and tok.endswith('**'):
                    r = p.add_run(tok[2:-2]); r.bold = True
                elif tok.startswith('`') and tok.endswith('`'):
                    r = p.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9)
                elif tok.startswith('$') and tok.endswith('$'):
                    r = p.add_run(tok); r.italic = True
                else:
                    p.add_run(tok)
            i += 1
            continue

        # 普通段落 (跳过空行)
        if stripped == '':
            i += 1
            continue

        p = doc.add_paragraph()
        tokens = re.split(r'(\*\*.+?\*\*|`[^`]+`|\$[^$]+\$)', stripped)
        for tok in tokens:
            if not tok:
                continue
            if tok.startswith('**') and tok.endswith('**'):
                r = p.add_run(tok[2:-2]); r.bold = True
            elif tok.startswith('`') and tok.endswith('`'):
                r = p.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9)
            elif tok.startswith('$') and tok.endswith('$'):
                r = p.add_run(tok); r.italic = True; r.font.name = 'Cambria Math'
            else:
                p.add_run(tok)
        i += 1

    doc.save(out_path)
    print(f'已生成 Word: {out_path}')


if __name__ == '__main__':
    base = os.path.dirname(os.path.abspath(__file__))
    md = os.path.join(base, '论文_杆系微元Token化结构地震响应代理模型.md')
    out = os.path.join(base, '论文_杆系微元Token化结构地震响应代理模型.docx')
    convert(md, out)
